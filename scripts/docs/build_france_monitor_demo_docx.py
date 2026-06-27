from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "demo" / "france-monitor-demo-candidatures.md"
OUTPUT = ROOT / "docs" / "demo" / "france-monitor-demo-candidatures.docx"


def set_spacing(style, before: int = 0, after: int = 8, line: float = 1.15) -> None:
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    set_spacing(normal, after=8, line=1.15)

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, RGBColor(0, 0, 0)),
        ("Heading 2", 16, 18, 6, RGBColor(0, 0, 0)),
        ("Heading 3", 14, 16, 4, RGBColor(67, 67, 67)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = color
        set_spacing(style, before=before, after=after, line=1.15)


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(67, 67, 67)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_numbered(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_paragraph_with_inline_code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(67, 67, 67)
        else:
            run = para.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(11)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "DADCE0")
        borders.append(tag)
    tbl_pr.append(borders)


def build_doc() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_page(doc)
    configure_styles(doc)

    in_code = False
    code_lines: list[str] = []
    first_heading = True

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(title)
            run.font.name = "Arial"
            run.font.size = Pt(26)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            first_heading = False
            continue

        if line.startswith("## "):
            if not first_heading and line.startswith("## 1. "):
                pass
            doc.add_heading(line[3:].strip(), level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue

        number_match = re.match(r"^([0-9]+)\\.\\s+(.*)$", line)
        if number_match:
            add_numbered(doc, number_match.group(2).strip())
            continue

        add_paragraph_with_inline_code(doc, line)

    if code_lines:
        add_code_block(doc, code_lines)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
